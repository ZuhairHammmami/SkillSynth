"""Build the bilingual RTL report DOCX from two Markdown chapters plus rendered figure assets.

Caller: repository gate (.venv/bin/python scripts/build_report_docx.py); callee: python-docx API
writing docs/report/SkillSynth-Report-AR.docx. Supports --ch1/--ch2/--assets/--out overrides for
isolated self-tests. Missing inputs or assets fail fast via SystemExit; success exits 0 after
printing paragraph/image counts.
"""

import argparse
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
CH1 = ROOT / "docs" / "report" / "chapter-1-technologies.md"
CH2 = ROOT / "docs" / "report" / "chapter-2-project.md"
ASSETS = ROOT / "docs" / "report" / "assets"
OUT = ROOT / "docs" / "report" / "SkillSynth-Report-AR.docx"

FIG_CAPTIONS = {
    "erd": "\u0627\u0644\u0634\u0643\u0644 1: \u0645\u062e\u0637\u0637 \u0627\u0644\u0643\u064a\u0627\u0646\u0627\u062a \u0648\u0627\u0644\u0639\u0644\u0627\u0642\u0627\u062a \u0644\u0642\u0627\u0639\u062f\u0629 \u0627\u0644\u0628\u064a\u0627\u0646\u0627\u062a",
    "usecase-student": "\u0627\u0644\u0634\u0643\u0644 2: \u0645\u062e\u0637\u0637 \u062d\u0627\u0644\u0627\u062a \u0627\u0644\u0627\u0633\u062a\u062e\u062f\u0627\u0645 \u0644\u0644\u0637\u0627\u0644\u0628",
    "usecase-admin": "\u0627\u0644\u0634\u0643\u0644 3: \u0645\u062e\u0637\u0637 \u062d\u0627\u0644\u0627\u062a \u0627\u0644\u0627\u0633\u062a\u062e\u062f\u0627\u0645 \u0644\u0644\u0645\u0634\u0631\u0641",
}

INLINE_RE = re.compile(r"(\*\*.+?\*\*|\*[^*\s][^*]*?\*|`[^`]+`)")
ARABIC_RE = re.compile(r"[\u0600-\u06FF]")
HEADING_SIZES = {1: 16, 2: 14, 3: 13}


def parse_args(argv):
    """Parse optional path overrides so the builder can run against self-test fixtures.

    Caller: main; Callee: argparse.
    """
    p = argparse.ArgumentParser(description="Build SkillSynth-Report-AR.docx")
    p.add_argument("--ch1", type=Path, default=CH1, help="chapter 1 markdown path")
    p.add_argument("--ch2", type=Path, default=CH2, help="chapter 2 markdown path")
    p.add_argument("--assets", type=Path, default=ASSETS, help="figure assets directory")
    p.add_argument("--out", type=Path, default=OUT, help="output docx path")
    return p.parse_args(argv)


def read_lines(path, label):
    """Read a chapter file as lines, exiting with a clear message when it is missing.

    Caller: main; Callee: none.
    """
    if not path.is_file():
        raise SystemExit(f"report-builder: missing {label} file: {path}")
    return path.read_text(encoding="utf-8").splitlines()


def iter_blocks(lines):
    """Yield (kind, payload) markdown blocks: h1/h2/h3/p/bullet/fig from limited MD syntax.

    Caller: build_document; Callee: none. Consecutive plain lines merge into one paragraph.
    """
    para = []
    for raw in lines:
        line = raw.rstrip()
        s = line.strip()
        if line.startswith("### "):
            if para:
                yield "p", " ".join(para)
                para = []
            yield "h3", line[4:].strip()
        elif line.startswith("## "):
            if para:
                yield "p", " ".join(para)
                para = []
            yield "h2", line[3:].strip()
        elif line.startswith("# "):
            if para:
                yield "p", " ".join(para)
                para = []
            yield "h1", line[2:].strip()
        elif line.startswith("- "):
            if para:
                yield "p", " ".join(para)
                para = []
            yield "bullet", line[2:].strip()
        elif s.startswith("[[FIG:") and s.endswith("]]"):
            if para:
                yield "p", " ".join(para)
                para = []
            yield "fig", s[6:-2].strip()
        elif not s:
            if para:
                yield "p", " ".join(para)
                para = []
        else:
            para.append(s)
    if para:
        yield "p", " ".join(para)


def set_rtl(run):
    """Mark one run as complex-script (Arabic) with matching szCs so Word renders it RTL.

    Caller: add_inline; Callee: none.
    """
    rpr = run._element.get_or_add_rPr()
    if rpr.find(qn("w:rtl")) is None:
        rpr.append(OxmlElement("w:rtl"))
    if run.font.size is not None:
        half = str(int(run.font.size.pt * 2))
        szcs = OxmlElement("w:szCs")
        szcs.set(qn("w:val"), half)
        rpr.append(szcs)


def style_run(run, bold=False, italic=False, mono=False):
    """Apply inline formatting to a run and flag it RTL only when it contains Arabic.

    Caller: add_inline, add_heading; Callee: set_rtl.
    """
    run.font.name = "Consolas" if mono else "Tajawal"
    run.font.bold = bold
    run.font.italic = italic
    if not mono:
        run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:cs"), "Tajawal")
    text = run.text or ""
    if ARABIC_RE.search(text):
        set_rtl(run)


def add_inline(par, text, base_bold=False):
    """Parse **bold**, *italic*, `code` segments and append styled runs to a paragraph.

    Caller: add_heading, add_body; Callee: style_run.
    """
    for token in INLINE_RE.split(text):
        if not token:
            continue
        if token.startswith("**") and token.endswith("**") and len(token) > 4:
            style_run(par.add_run(token[2:-2]), bold=True)
        elif token.startswith("`") and token.endswith("`") and len(token) > 2:
            style_run(par.add_run(token[1:-1]), bold=base_bold, mono=True)
        elif token.startswith("*") and token.endswith("*") and len(token) > 2:
            style_run(par.add_run(token[1:-1]), italic=True)
        else:
            style_run(par.add_run(token), bold=base_bold)


def para_bidi(par, align):
    """Force right-to-left paragraph direction (w:bidi) and set the requested alignment.

    Caller: add_heading, add_body, add_figure; Callee: none.
    """
    par.paragraph_format.alignment = align
    ppr = par._p.get_or_add_pPr()
    if ppr.find(qn("w:bidi")) is None:
        ppr.append(OxmlElement("w:bidi"))


def strip_md(text):
    """Remove inline markdown markers from heading text before rendering.

    Caller: add_heading; Callee: none.
    """
    out = INLINE_RE.sub(lambda m: m.group(0).strip("*`"), text)
    return out.replace("**", "").replace("`", "")


def add_heading(doc, text, level):
    """Append a Tajawal heading (16/14/13pt) using real Heading styles for navigation.

    Caller: build_document; Callee: strip_md, style_run, para_bidi.
    """
    par = doc.add_paragraph(style=f"Heading {level}")
    run = par.add_run(strip_md(text))
    run.font.size = Pt(HEADING_SIZES[level])
    style_run(run, bold=True)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    par.paragraph_format.space_before = Pt(10 if level == 1 else 8)
    par.paragraph_format.space_after = Pt(6)
    para_bidi(par, WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.RIGHT)


def add_body(doc, text, bullet=False):
    """Append an RTL body or bullet paragraph carrying parsed inline runs.

    Caller: build_document; Callee: add_inline, para_bidi.
    """
    par = doc.add_paragraph(style="List Bullet" if bullet else None)
    add_inline(par, text)
    para_bidi(par, WD_ALIGN_PARAGRAPH.RIGHT)
    return par


def add_figure(doc, assets_dir, name):
    """Append a centered 6-inch figure with its mapped Arabic caption below it.

    Caller: build_document; Callee: style_run, para_bidi.
    """
    if name not in FIG_CAPTIONS:
        raise SystemExit(f"report-builder: no caption mapping for [[FIG:{name}]]")
    path = assets_dir / f"{name}.png"
    if not path.is_file():
        raise SystemExit(f"report-builder: missing asset image: {path}")
    pic = doc.add_paragraph()
    pic.add_run().add_picture(str(path), width=Inches(6.0))
    para_bidi(pic, WD_ALIGN_PARAGRAPH.CENTER)
    cap = doc.add_paragraph()
    style_run(cap.add_run(FIG_CAPTIONS[name]))
    cap.paragraph_format.space_after = Pt(10)
    para_bidi(cap, WD_ALIGN_PARAGRAPH.CENTER)


def setup_styles(doc):
    """Configure the Normal style: Tajawal 12pt (with complex-script font), 1.15 spacing.

    Caller: build_document; Callee: none.
    """
    normal = doc.styles["Normal"]
    normal.font.name = "Tajawal"
    normal.font.size = Pt(12)
    normal.element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:cs"), "Tajawal")
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)
    bullet_style = doc.styles["List Bullet"]
    bullet_style.base_style = normal


def build_document(blocks, assets_dir):
    """Fill a new Document from parsed blocks, dispatching each block kind to its renderer.

    Caller: main; Callee: setup_styles, add_heading, add_body, add_figure, iter_blocks.
    """
    doc = Document()
    setup_styles(doc)
    for kind, payload in blocks:
        if kind == "fig":
            add_figure(doc, assets_dir, payload)
        elif kind.startswith("h"):
            add_heading(doc, payload, int(kind[1]))
        elif kind == "bullet":
            add_body(doc, payload, bullet=True)
        else:
            add_body(doc, payload)
    return doc


def main(argv=None):
    """Read both chapters, build the DOCX, save it, and print paragraph/image statistics.

    Caller: CLI (__main__); Callee: parse_args, read_lines, iter_blocks, build_document.
    """
    args = parse_args(sys.argv[1:] if argv is None else argv)
    blocks = []
    for label, path in (("chapter-1", args.ch1), ("chapter-2", args.ch2)):
        lines = read_lines(path, label)
        blocks.extend(iter_blocks(lines))
    doc = build_document(blocks, args.assets)
    out_parent = args.out.parent
    if not out_parent.is_dir():
        raise SystemExit(f"report-builder: output directory does not exist: {out_parent}")
    doc.save(str(args.out))
    print(f"[report] wrote {args.out}")
    print(f"[report] paragraphs={len(doc.paragraphs)} images={len(doc.inline_shapes)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
